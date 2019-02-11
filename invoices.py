# Invoice Generator
# by Nathan Lam (nathan@codecreate.com.au)
# created on 25/12/2018

from docx import *
import time, os

def main():
    invoice = Invoice()
    name = input('Name: ')
    abn = input('ABN: ')
    acc_name = input('Bank account name: ')
    bsb = input('BSB: ')
    acc_no = input('Account number: ')
    addr = input('From address, first line: ')
    city = input('City: ')
    state = input('State: ')
    postcode = input('Postcode: ')
    phone = input('Mobile number: ')

    invoice.set_name(name)
    invoice.set_abn(abn)
    invoice.set_bank_details(acc_name, bsb, acc_no)
    invoice.set_from_details(addr, city, state, postcode, phone)

    hasMoreItems = True
    while hasMoreItems:

        user_input = input('Do you want to insert a new item (Y/N)? ')
        choice = user_input.upper()
        if choice == 'N':
            hasMoreItems = False
        elif choice == 'Y':
            title = input('Item title: ')
            desc = input('Item description: ')
            date = input('Date: ')
            try:
                amount = float(input('Amount owed (amount without dollar sign): '))
                invoice.insert_item(title, desc, date, amount)
            except:
                print('Invalid input: amount should only be a float')
        else:
            print('Invalid input')


class Invoice:
    def __init__(self, template_name='InvoiceTemplate.docx'):
        try:
            template = os.path.join(os.path.dirname(__file__), template_name)
            self.document = Document(template)
            self.max_items = len(self.document.tables[2].rows) - 1
            invoice_no = time.strftime('%Y%m%d')
            invoice_no_field = self.document.tables[0].rows[1].cells[1].paragraphs[0].runs[4]
            invoice_no_field.text = invoice_no
            invoice_date = time.strftime('%B %d, %Y').upper()
            date_field = self.document.tables[0].rows[1].cells[1].paragraphs[1].runs[2]
            date_field.text = invoice_date
        except:
            print('No template document available')
        self.items = 0
        self.total_amount = 0

    def __del__(self):
        self.document.save('NewInvoice.docx')

    def insert_item(self, title='', description='', date='', amount=0):
        if self.items >= self.max_items:
            return
        rows = self.document.tables[2].rows[self.items + 1]
        rows.cells[0].text = title
        rows.cells[1].text = description
        rows.cells[2].text = date
        rows.cells[3].text = '$' + f'{amount:.2f}'
        self.items += 1
        self.total_amount += amount
        total_field = self.document.tables[2].rows[self.max_items].cells[3]
        total_field.text = '$' + f'{self.total_amount:.2f}'

    def set_name(self, name):
        if type(name) is not str:
            return
        name_field = self.document.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
        name_field.text = name.upper()


    def set_abn(self, abn):
        if type(abn) is not str:
            print('Invalid input, enter your ABN as a string')
            return
        abn_field = self.document.tables[0].rows[0].cells[0].paragraphs[1].runs[1]
        abn_field.text = abn

    def set_from_details(self, address, city, state, postcode, phone):
        addr_field = self.document.tables[0].rows[0].cells[0].paragraphs[2].runs[0]
        addr_field.text = address
        city_field = self.document.tables[0].rows[0].cells[0].paragraphs[2].runs[2]
        city_field.text = city.capitalize()
        state_field = self.document.tables[0].rows[0].cells[0].paragraphs[2].runs[3]
        state_field = state.upper()
        postcode_field = self.document.tables[0].rows[0].cells[0].paragraphs[2].runs[4]
        postcode_field = str(postcode)
        phone_field = self.document.tables[0].rows[0].cells[0].paragraphs[3].runs[3]
        phone_field = str(phone)

    def set_bank_details(self, acc_name, bsb, acc_no):
        if type(acc_name) is not str:
            print('Invalid parameter: all parameters should be strings')
        acc_name_field = self.document.paragraphs[2].runs[2]
        acc_name_field.text = acc_name.upper()
        bsb_field = self.document.paragraphs[2].runs[5]
        bsb_field.text = str(bsb)
        acc_no_field = self.document.paragraphs[2].runs[8]
        acc_no_field.text = str(acc_no)

    def set_invoice_purpose(self, purpose):
        if type(purpose) is not str:
            print('Invalid parameter: all parameters should be strings')
        invoice_purpose_field = self.document.tables[1].rows[0].cell[1].paragraphs[1].runs[0]
        invoice_purpose_field.text = purpose

    def show_doc(self):
        i = 0
        for p in self.document.paragraphs:
            print('Paragraph', i)
            i += 1
            j = 0
            for r in p.runs:
                print('\tRun', j, '->', r.text)
                j += 1

        i = 0
        for s in self.document.sections:
            print('Section', i, '->', s)
            i += 1

        i = 0
        for t in self.document.tables[0:2]:
            print('Table', i, '->', t)
            i += 1
            j = 0
            for r in t.rows:
                print('\tRow', j, '->', r)
                j += 1
                k = 0
                for c in r.cells:
                    print('\t\tCell', k, '->', c)
                    k += 1
                    l = 0
                    for p in c.paragraphs:
                        print('\t\t\tParagraph', l)
                        l += 1
                        m = 0
                        for r in p.runs:
                            print('\t\t\t\tRun', m, '->', r.text)
                            m += 1

if __name__ == '__main__':
    main()
